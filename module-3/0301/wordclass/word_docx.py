#!/usr/bin/env python
#operation for word
#author:wuchongyao
#time:2018-08-16


import docx


class test_docx():
	# def __init__(self,):
		# pass

	@staticmethod
	def test_title(content, title_level, dst_name, src_name):
		'''将输入内容保存为标题'''
		document = docx.Document(src_name)
		document.add_heading(content, title_level)
		document.save(dst_name)
		

	@staticmethod
	def test_subtitle(content,dst_name, src_name):
		'''将输入内容保存为副标题'''
		document = docx.Document(src_name)
		document.add_paragraph(content, 'Subtitle')
		document.save(dst_name)
		
		
	@staticmethod
	def test_paragraph(content, dst_name, src_name):
		'''将输入内容保存为正文'''
		document = docx.Document(src_name)
		document.add_paragraph(content)
		document.save(dst_name)
		

def main():
	'''测试代码'''
	title = "热爱生命"
	subtitle = "——汪国真"
	content = '''
我不去想是否能够成功
既然选择了远方
便只顾风雨兼程

我不去想能否赢得爱情
既然钟情于玫瑰
就勇敢地吐露真诚

我不去想身后会不会袭来寒风冷雨
既然目标是地平线
留给世界的只能是背影

我不去想未来是平坦还是泥泞
只要热爱生命
一切，都在意料之中
	'''
	title_level = 0
	dst_name = r"d:/python/practice/热爱生命.docx"
	src_name = r"d:/python/practice/热爱生命.docx"
	test_docx.test_title(title, title_level, dst_name, None)
	test_docx.test_subtitle(subtitle, dst_name, src_name)
	test_docx.test_paragraph(content, dst_name, src_name)


if __name__ == "__main__":
	main()
		

